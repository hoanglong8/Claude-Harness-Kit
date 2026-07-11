#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
foxai-book-writer v4 — generate_book.py
Design: Navy/Gold premium (inspired by The AI Operator series)

New features vs v3:
  • Navy + Gold color scheme (was dark-brown)
  • Thin gold border frame on every page
  • No running header — minimal footer (series label left + page number right)
  • Chapter heading drawn on first content page (navy bold + gold underline)
  • TOC redesigned: large gray chapter numbers + gold separator rules
  • Callout boxes (!! prefix in content markdown)
  • DOCX: page border via <w:pgBorders> XML
  • JSON new fields: series_label, tags, level_badge

Usage:
    python generate_book.py <book_data.json> [output_dir]

New JSON fields (all optional):
    book.series_label   → "THE AI OPERATOR · TẬP 2"
    book.tags           → ["CLAUDE COWORK", "CLAUDE CODE", "CODEX"]
    book.level_badge    → "ADVANCED · CẤP ĐỘ NÂNG CAO"
    chapter.content     → !! prefix = callout box
"""

import json, sys, os, re
from pathlib import Path
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm, mm
from reportlab.lib.colors import HexColor, Color, white, black
from reportlab.platypus import Paragraph, Frame, HRFlowable, Spacer, Table, TableStyle
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT
from reportlab.pdfgen import canvas as pdf_canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.pdfmetrics import registerFontFamily

from docx import Document
from docx.shared import Pt, Cm as DCm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ── Paths ─────────────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).parent
ASSETS_DIR = SCRIPT_DIR.parent / 'assets'
FONTS_DIR  = ASSETS_DIR / 'fonts'
IMAGES_DIR = ASSETS_DIR / 'images'

# ── Themes ─────────────────────────────────────────────────────────────────────
# Chọn qua trường "theme" trong book JSON: "navy-gold" (mặc định) | "burgundy".
# Palette burgundy trích bằng đo pixel từ The AI Operator tập 1 (PDF gốc):
# bìa #9E1A31, heading trên nền cream #6E1121, accent #A8853B, nền #FBFBF8.
THEMES = {
    'navy-gold': dict(
        cover='#1E2A45', heading='#1E2A45', accent='#C9A84C',
        page_bg='#FDFCF8', callout_bg='#EEECE4', overlay=(0.12, 0.16, 0.27),
    ),
    'burgundy': dict(
        cover='#9E1A31', heading='#6E1121', accent='#A8853B',
        page_bg='#FBFBF8', callout_bg='#F4F2EB', overlay=(0.28, 0.07, 0.11),
    ),
}

def _peek_theme():
    """Đọc trường 'theme' từ book JSON (sys.argv[1]) TRƯỚC khi định nghĩa token —
    các hằng số bên dưới được dùng làm default argument nên phải chốt ngay lúc import."""
    try:
        with open(sys.argv[1], 'r', encoding='utf-8') as f:
            name = (json.load(f).get('theme') or 'navy-gold').strip().lower()
        if name not in THEMES:
            print(f"[theme] '{name}' khong ton tai — dung navy-gold. Co: {', '.join(THEMES)}")
            return 'navy-gold'
        return name
    except Exception:
        return 'navy-gold'

THEME_NAME = _peek_theme()
_T = THEMES[THEME_NAME]

# ── Design tokens (theo theme) ─────────────────────────────────────────────────
# Tên biến NAVY/GOLD giữ nguyên để tương thích toàn bộ code phía dưới.
COVER_BG    = HexColor(_T['cover'])     # Nền bìa + splash pages
NAVY        = HexColor(_T['heading'])   # Headings trên nền cream
GOLD        = HexColor(_T['accent'])    # Borders, accents, rules
PAGE_BG     = HexColor(_T['page_bg'])   # Body page background
BODY_TXT    = HexColor('#1A1A1A')   # Body text (near-black)
SUBTEXT     = HexColor('#555555')   # Subtitles, secondary
LIGHT_GRAY  = HexColor('#8A8A8A')   # Footer, small labels
TOC_GRAY    = HexColor('#CCCCCC')   # Large TOC numbers (decorative)
CALLOUT_BG  = HexColor(_T['callout_bg'])# Callout box background
DARK_OVERLAY= Color(*_T['overlay'], alpha=0.72)

BORDER_INS  = 0.60 * cm
BORDER_LW   = 0.8

# ── Labeled callout box definitions ───────────────────────────────────────────
# Key = markdown prefix (e.g. '!b'), Value = (label, bg_color, label_color)
CALLOUT_DEFS = {
    '!b': ('BỐI CẢNH',        HexColor('#EBF0F8'), HexColor('#1E3C6E')),
    '!t': ('THÀNH PHẨM',      HexColor('#EBF5EB'), HexColor('#1A5C2A')),
    '!d': ('ĐÒN BẨY',         HexColor('#FDF8EC'), HexColor('#8B6914')),
    '!l': ('LỖI THƯỜNG GẶP',  HexColor('#FDF0EC'), HexColor('#C0391E')),
    '!p': ('MẸO PRO',         HexColor('#F3ECFD'), HexColor('#6A1FAC')),
    '!c': ('⚙ CÁCH DỰNG',    HexColor('#F0F4F8'), HexColor('#1E3A5A')),
    '!g': ('✦ BẢN GIAO VIỆC', COVER_BG, GOLD),
}
# DOCX fill colors for labeled callouts (hex string without #)
CALLOUT_DOCX_FILLS = {
    '!b': 'EBF0F8', '!t': 'EBF5EB', '!d': 'FDF8EC',
    '!l': 'FDF0EC', '!p': 'F3ECFD', '!c': 'F0F4F8', '!g': _T['cover'][1:],
}
CALLOUT_DOCX_LABEL_COLORS = {
    '!b': '1E3C6E', '!t': '1A5C2A', '!d': '8B6914',
    '!l': 'C0391E', '!p': '6A1FAC', '!c': '1E3A5A', '!g': _T['accent'][1:],
}

# DOCX RGB
NAVY_RGB    = RGBColor(*bytes.fromhex(_T['heading'][1:]))
COVER_RGB   = RGBColor(*bytes.fromhex(_T['cover'][1:]))
GOLD_RGB    = RGBColor(*bytes.fromhex(_T['accent'][1:]))
BODY_RGB    = RGBColor(0x1A, 0x1A, 0x1A)
SUBTEXT_RGB = RGBColor(0x55, 0x55, 0x55)
LIGHT_RGB   = RGBColor(0x8A, 0x8A, 0x8A)
WHITE_RGB   = RGBColor(0xFF, 0xFF, 0xFF)
CREAM_RGB   = RGBColor(0xED, 0xE0, 0xC8)

# Page layout
PAGE_W, PAGE_H = A4
ML = 2.5*cm; MR = 2.0*cm
MT = 2.4*cm; MB = 2.4*cm
CW = PAGE_W - ML - MR
CH = PAGE_H - MT - MB
HEADING_H = 4.0*cm   # space reserved for chapter heading on first content page

_ORDINALS = {
    1:'MỘT',2:'HAI',3:'BA',4:'BỐN',5:'NĂM',6:'SÁU',7:'BẢY',8:'TÁM',
    9:'CHÍN',10:'MƯỜI',11:'MƯỜI MỘT',12:'MƯỜI HAI',13:'MƯỜI BA',
    14:'MƯỜI BỐN',15:'MƯỜI LĂM',
}
def ordinal(n): return _ORDINALS.get(n, str(n))
def slug(s, n=35):
    return '_'.join(re.sub(r'[^\w ]','',s,flags=re.UNICODE).split())[:n]

# ── Font management ────────────────────────────────────────────────────────────
FD = 'CG'; FB = 'BookTNR'
_FONT_DONE = False; _FD_OVERRIDE = None

def _dl_fonts():
    import urllib.request, re as _re
    FONTS_DIR.mkdir(parents=True, exist_ok=True)
    variants = [(0,400,'CormorantGaramond-Regular'),(0,700,'CormorantGaramond-Bold'),
                (1,400,'CormorantGaramond-Italic'),(1,700,'CormorantGaramond-BoldItalic')]
    UA = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) Chrome/120.0.0.0'
    for ital, wght, base in variants:
        dst = FONTS_DIR / f'{base}.ttf'
        if dst.exists(): continue
        url = (f'https://fonts.googleapis.com/css2?family=Cormorant+Garamond'
               f':ital,wght@{ital},{wght}&display=swap')
        print(f'  Downloading {base}...', end='', flush=True)
        try:
            req = urllib.request.Request(url, headers={'User-Agent': UA})
            with urllib.request.urlopen(req, timeout=15) as r: css = r.read().decode()
            urls = _re.findall(r'url\((https://fonts\.gstatic\.com[^)]+)\)', css)
            if not urls: print(' no URL'); continue
            req2 = urllib.request.Request(urls[0].strip("'\""), headers={'User-Agent': UA})
            with urllib.request.urlopen(req2, timeout=30) as r2: data = r2.read()
            magic = int.from_bytes(data[:4], 'big')
            if magic == 0x00010000:
                dst.write_bytes(data); print(f' {len(data)//1024}KB ok')
            elif magic == 0x774F4632:
                print(f' WOFF2', end='')
                try:
                    from fontTools.ttLib import TTFont as FT
                    import io; ft = FT(io.BytesIO(data)); ft.flavor = None; ft.save(str(dst))
                    print(' → TTF ok')
                except Exception as e: print(f' ({e})')
            else: print(f' fmt 0x{magic:08x}')
        except Exception as e: print(f' SKIP ({e})')

def _reg_fonts():
    global _FONT_DONE, _FD_OVERRIDE
    if _FONT_DONE: return
    _FONT_DONE = True
    cg = FONTS_DIR / 'CormorantGaramond-Regular.ttf'
    if not cg.exists():
        print('  Setting up Cormorant Garamond (first run)...'); _dl_fonts()
    cg_map = [(FD,'CormorantGaramond-Regular.ttf'),(FD+'-Bold','CormorantGaramond-Bold.ttf'),
              (FD+'-Italic','CormorantGaramond-Italic.ttf'),(FD+'-BoldItalic','CormorantGaramond-BoldItalic.ttf')]
    ok = True
    for name, fname in cg_map:
        p = FONTS_DIR / fname
        if p.exists():
            try: pdfmetrics.registerFont(TTFont(name, str(p)))
            except: ok = False
        else: ok = False
    if ok:
        registerFontFamily(FD, normal=FD, bold=FD+'-Bold', italic=FD+'-Italic', boldItalic=FD+'-BoldItalic')
        print('  Display font: Cormorant Garamond registered')
    else:
        print('  Display font: fallback to Times New Roman')
        _FD_OVERRIDE = FB
    tnr_dirs = [r'C:\Windows\Fonts', '/usr/share/fonts/truetype/msttcorefonts', '/Library/Fonts']
    tnr_files = [('times.ttf',FB),('timesbd.ttf',FB+'-Bold'),('timesi.ttf',FB+'-Italic'),('timesbi.ttf',FB+'-BoldItalic')]
    for d in tnr_dirs:
        paths = [(os.path.join(d,f),n) for f,n in tnr_files]
        if not all(os.path.exists(p) for p,_ in paths): continue
        try:
            for p,n in paths: pdfmetrics.registerFont(TTFont(n,p))
            registerFontFamily(FB, normal=FB, bold=FB+'-Bold', italic=FB+'-Italic', boldItalic=FB+'-BoldItalic')
            print(f'  Body font: Times New Roman registered from {d}'); break
        except Exception as e: print(f'  TNR: {e}')

def FDN(bold=False, italic=False):
    if not _FONT_DONE: _reg_fonts()
    f = _FD_OVERRIDE or FD
    if bold and italic: return f+'-BoldItalic'
    if bold: return f+'-Bold'
    if italic: return f+'-Italic'
    return f

def FBN(bold=False, italic=False):
    if not _FONT_DONE: _reg_fonts()
    if bold and italic: return FB+'-BoldItalic'
    if bold: return FB+'-Bold'
    if italic: return FB+'-Italic'
    return FB

# ── Canvas helpers ─────────────────────────────────────────────────────────────
def _txt(c, s, x, y, font, size, color, align='left'):
    c.saveState(); c.setFont(font, size); c.setFillColor(color)
    if   align=='center': c.drawCentredString(x, y, s)
    elif align=='right':  c.drawRightString(x, y, s)
    else:                 c.drawString(x, y, s)
    c.restoreState()

def _wrap(c, s, x, y, w, font, size, color, align='left', lh=None):
    if not s: return y
    if lh is None: lh = size * 1.45
    c.setFont(font, size)
    words = s.split(); lines = []; cur = []
    for word in words:
        test = ' '.join(cur+[word])
        if c.stringWidth(test, font, size) <= w: cur.append(word)
        else:
            if cur: lines.append(' '.join(cur))
            cur = [word]
    if cur: lines.append(' '.join(cur))
    for line in lines:
        _txt(c, line, x+(w/2 if align=='center' else 0), y, font, size, color, align)
        y -= lh
    return y

def _hline(c, x1, y, x2, w=0.5, color=GOLD):
    c.saveState(); c.setStrokeColor(color); c.setLineWidth(w)
    c.line(x1, y, x2, y); c.restoreState()

def _page_frame(c, color=GOLD, inset=BORDER_INS, lw=BORDER_LW):
    """Thin gold border frame on every page."""
    c.saveState(); c.setStrokeColor(color); c.setLineWidth(lw)
    c.rect(inset, inset, PAGE_W-2*inset, PAGE_H-2*inset, fill=0, stroke=1)
    c.restoreState()

def _bg_navy(c, image=None):
    """Navy background (cover, splash pages) with optional image overlay."""
    c.setFillColor(COVER_BG); c.rect(0,0,PAGE_W,PAGE_H,fill=1,stroke=0)
    if image and os.path.exists(image):
        try:
            c.drawImage(image,0,0,PAGE_W,PAGE_H,preserveAspectRatio=False,mask='auto')
            c.saveState(); c.setFillColor(DARK_OVERLAY)
            c.rect(0,0,PAGE_W,PAGE_H,fill=1,stroke=0); c.restoreState()
        except: pass

def _bg_cream(c):
    c.setFillColor(PAGE_BG); c.rect(0,0,PAGE_W,PAGE_H,fill=1,stroke=0)

def _lspaced(text, spacing=0.12):
    """Approximate letter-spacing by inserting thin spaces (visual only)."""
    return (' ' * 1).join(list(text))  # simple version

# ── PDF page builders ──────────────────────────────────────────────────────────
def _p_cover(c, book):
    """Navy cover: border frame + series label + big title + tags + author + badge."""
    _bg_navy(c, book.get('cover_image'))
    _page_frame(c, color=GOLD, inset=BORDER_INS+0.05*cm, lw=1.2)

    # Series label (top, small caps gold, letter-spaced)
    series = book.get('series_label', '')
    if series:
        _txt(c, series.upper(), PAGE_W/2, PAGE_H-1.85*cm, FDN(), 8, GOLD, 'center')

    # Title (large, white, centered, ~38-52% from top)
    title_words = book['title'].upper()
    ty = PAGE_H * 0.52
    c.setFont(FDN(bold=True), 52)
    tw = c.stringWidth(title_words, FDN(bold=True), 52)
    avail_w = CW * 0.85
    if tw <= avail_w:
        _txt(c, title_words, PAGE_W/2, ty, FDN(bold=True), 52, white, 'center')
        ty -= 52*0.04*cm + 0.3*cm
    else:
        # Wrap at ~36pt if title is long
        ty = _wrap(c, title_words, PAGE_W*0.08, ty, PAGE_W*0.84,
                   FDN(bold=True), 40, white, 'center', lh=46)
        ty -= 0.3*cm

    # Gold rule
    _hline(c, PAGE_W*0.34, ty, PAGE_W*0.66, w=1.8)
    ty -= 1.0*cm

    # Subtitle (italic, cream)
    if book.get('subtitle'):
        ty = _wrap(c, book['subtitle'], PAGE_W*0.10, ty, PAGE_W*0.80,
                   FDN(italic=True), 11.5, HexColor('#E8DFC8'), 'center', lh=16)
        ty -= 0.35*cm

    # Quote (small italic gold)
    if book.get('cover_quote'):
        ty = _wrap(c, f'“{book["cover_quote"]}”',
                   PAGE_W*0.12, ty, PAGE_W*0.76,
                   FBN(italic=True), 9.5, GOLD, 'center', lh=13)

    # Tags row (small caps, letter-spaced, light)
    tags = book.get('tags', [])
    tag_y = PAGE_H * 0.18
    if tags:
        tag_str = '  ·  '.join(t.upper() for t in tags)
        _txt(c, tag_str, PAGE_W/2, tag_y, FBN(), 8, HexColor('#B0A898'), 'center')
        tag_y -= 0.65*cm

    # Author
    _txt(c, book.get('author','').upper() if book.get('author') else '',
         PAGE_W/2, tag_y, FDN(bold=True), 14, white, 'center')
    if book.get('author_byline') or book.get('coauthor'):
        byline = book.get('author_byline') or book.get('coauthor', '')
        _txt(c, byline.upper(), PAGE_W/2, tag_y-0.6*cm, FBN(), 7.5,
             HexColor('#7A8A9A'), 'center')

    # Level badge (outlined box at very bottom)
    badge = book.get('level_badge', '')
    if badge:
        bw = c.stringWidth(badge.upper(), FBN(), 8) + 1.2*cm
        bx = PAGE_W/2 - bw/2; by = MB + 0.6*cm
        c.saveState()
        c.setStrokeColor(GOLD); c.setLineWidth(0.7)
        c.rect(bx, by, bw, 0.55*cm, fill=0, stroke=1)
        _txt(c, badge.upper(), PAGE_W/2, by+0.15*cm, FBN(), 8, GOLD, 'center')
        c.restoreState()

def _setup_content_page(c, pn, series_label=''):
    """Cream bg + gold border + minimal footer. Call at start of every body page."""
    _bg_cream(c)
    _page_frame(c)
    # Footer
    _hline(c, ML, MB-0.3*cm, PAGE_W-MR, w=0.35, color=GOLD)
    _txt(c, series_label.upper() if series_label else '',
         ML, MB-0.7*cm, FBN(), 7, LIGHT_GRAY)
    _txt(c, str(pn), PAGE_W-MR, MB-0.7*cm, FBN(), 8, LIGHT_GRAY, 'right')

def _draw_chapter_heading(c, ch):
    """Draw chapter heading at top of first content page. Navy bold + gold rule."""
    y = PAGE_H - MT - 0.15*cm
    num_str = f'{ch["number"]:02d}' if isinstance(ch.get('number'), int) else str(ch.get('number',''))
    lbl = f"CHƯƠNG  {num_str}"
    _txt(c, lbl, ML, y, FDN(italic=True), 9, GOLD)
    y -= 0.75*cm
    y = _wrap(c, ch['title'], ML, y, CW, FDN(bold=True), 22, NAVY, lh=26)
    y -= 0.4*cm
    _hline(c, ML, y, ML + 3.8*cm, w=2.2, color=GOLD)
    y -= 0.35*cm
    if ch.get('subtitle'):
        _txt(c, ch['subtitle'], ML, y, FDN(italic=True), 10.5, SUBTEXT)

def _p_about(c, book, pn):
    """About / preface page — content style with gold heading."""
    series = book.get('series_label', book.get('title', ''))
    _setup_content_page(c, pn, series)

    y = PAGE_H - MT - 0.15*cm
    about_title = book.get('about_title', 'VỀ TÁC GIẢ')
    _txt(c, about_title, ML, y, FDN(bold=True), 24, NAVY)
    y -= 0.55*cm
    _hline(c, ML, y, ML + 3.5*cm, w=2.2, color=GOLD)
    y -= 0.8*cm

    if book.get('about_subtitle'):
        y = _wrap(c, book['about_subtitle'], ML, y, CW,
                  FDN(italic=True), 11, SUBTEXT, lh=15)
        y -= 0.5*cm

    if book.get('author'):
        _txt(c, book['author'], ML, y, FDN(bold=True), 13, GOLD)
        y -= 0.5*cm
    if book.get('author_title'):
        _txt(c, book['author_title'], ML, y, FBN(italic=True), 9.5, LIGHT_GRAY)
        y -= 0.55*cm
    if book.get('author_contact'):
        _txt(c, book['author_contact'], ML, y, FBN(), 9, LIGHT_GRAY)
        y -= 0.7*cm

    if book.get('about_author'):
        _hline(c, ML, y+0.15*cm, PAGE_W-MR, w=0.3)
        y -= 0.5*cm
        y = _wrap(c, book['about_author'], ML, y, CW, FBN(), 10.5, BODY_TXT, lh=15)
        y -= 0.5*cm

    if book.get('credits'):
        _hline(c, ML, y+0.1*cm, PAGE_W-MR, w=0.3)
        y -= 0.5*cm
        creds = list(book['credits'].items())
        cw = CW / len(creds)
        for i, (k, v) in enumerate(creds):
            cx = ML + cw*i + cw/2
            _txt(c, k.upper(), cx, y, FBN(), 7.5, LIGHT_GRAY, 'center')
            _txt(c, v, cx, y-14, FDN(bold=True), 9.5, GOLD, 'center')

def _p_toc(c, book, pn):
    """TOC with large gray chapter numbers + gold separator rules."""
    series = book.get('series_label', book.get('title',''))
    _setup_content_page(c, pn, series)

    # Heading
    y = PAGE_H - MT - 0.15*cm
    _txt(c, 'MỤC LỤC', ML, y, FDN(bold=True), 24, NAVY)
    y -= 0.5*cm
    _hline(c, ML, y, ML+3.5*cm, w=2.2, color=GOLD)
    y -= 0.7*cm

    num_x  = ML
    title_x = ML + 1.9*cm
    title_w = CW - 1.9*cm - 1.0*cm
    pg_x   = PAGE_W - MR

    EST_PG = 9  # starting page estimate
    for part in book.get('parts', []):
        # Part label
        part_lbl = f"PHẦN {part['number']}  ·  {part.get('title','').upper()}"
        _txt(c, part_lbl, num_x, y, FDN(bold=True), 8.5, GOLD)
        y -= 0.65*cm

        for ch in part.get('chapters', []):
            if y < MB + 2.0*cm: break  # safety: don't overflow
            # Large gray chapter number
            _txt(c, str(ch['number']), num_x+0.1*cm, y-0.1*cm, FDN(bold=True), 38, TOC_GRAY)

            # Title
            title_str = ch['title']
            c.setFont(FDN(bold=True), 11.5)
            tw = c.stringWidth(title_str, FDN(bold=True), 11.5)
            if tw <= title_w:
                _txt(c, title_str, title_x, y, FDN(bold=True), 11.5, NAVY)
            else:
                _wrap(c, title_str, title_x, y, title_w, FDN(bold=True), 10.5, NAVY, lh=14)

            # Subtitle
            if ch.get('subtitle'):
                _txt(c, ch['subtitle'], title_x, y-0.48*cm, FBN(italic=True), 8.5, LIGHT_GRAY)

            # Page number
            _txt(c, str(EST_PG), pg_x, y-0.1*cm, FBN(), 9, LIGHT_GRAY, 'right')
            EST_PG += 15

            # Gold separator
            y -= 1.9*cm
            _hline(c, ML, y+0.5*cm, PAGE_W-MR, w=0.3, color=GOLD)
            y -= 0.3*cm

        y -= 0.25*cm  # extra gap between parts

def _p_part_splash(c, part):
    """Navy part divider."""
    _bg_navy(c, part.get('image'))
    _page_frame(c, color=GOLD, lw=1.1)
    cy = PAGE_H * 0.50
    _txt(c, f"PHẦN {str(part['number']).upper()}", PAGE_W/2,
         cy+1.7*cm, FDN(italic=True), 10, GOLD, 'center')
    _hline(c, PAGE_W*0.30, cy+0.95*cm, PAGE_W*0.70, w=1.0)
    _wrap(c, part.get('title','').upper(), PAGE_W*0.08, cy, PAGE_W*0.84,
          FDN(bold=True), 30, white, 'center', lh=35)
    if part.get('subtitle'):
        _txt(c, part['subtitle'], PAGE_W/2, cy-1.4*cm, FDN(italic=True), 12,
             HexColor('#D0C8B8'), 'center')

def _p_chapter_opener(c, ch):
    """Navy chapter opener splash — pure programmatic, no image required."""
    _bg_navy(c)   # pure navy; image overlay is optional via ch.get('image')
    if ch.get('image') and os.path.exists(ch['image']):
        try:
            c.drawImage(ch['image'],0,0,PAGE_W,PAGE_H,
                        preserveAspectRatio=False,mask='auto')
            c.saveState(); c.setFillColor(DARK_OVERLAY)
            c.rect(0,0,PAGE_W,PAGE_H,fill=1,stroke=0); c.restoreState()
        except: pass
    _page_frame(c, color=GOLD, lw=1.1)
    # "CHƯƠNG" small label at top (no ordinal word — matches AI Operator style)
    _txt(c, 'CHƯƠNG', ML, PAGE_H-MT, FDN(italic=True), 9, GOLD)
    # Large zero-padded number: "01", "02" …
    num_str = f'{ch["number"]:02d}' if isinstance(ch.get('number'), int) else str(ch.get('number',''))
    num_y = PAGE_H * 0.38
    _txt(c, num_str, ML, num_y, FDN(bold=True), 96, GOLD)
    title_y = num_y - 0.8*cm
    title_y = _wrap(c, ch['title'], ML, title_y, CW, FDN(bold=True), 26, white, lh=31)
    if ch.get('subtitle'):
        _txt(c, ch['subtitle'], ML, title_y-0.3*cm, FDN(italic=True), 12,
             HexColor('#C0B8A8'))
    if ch.get('quote'):
        _hline(c, ML, MB+4.5*cm, PAGE_W-MR, w=0.35)
        _wrap(c, f'“{ch["quote"]}”', ML, MB+4.2*cm, CW,
              FBN(italic=True), 9.5, HexColor('#D4CDB8'), lh=13)
        if ch.get('quote_attribution'):
            _txt(c, '— '+ch['quote_attribution'],
                 PAGE_W-MR, MB+1.0*cm, FDN(bold=True), 8, GOLD, 'right')

def _xe(s): return s.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
def _rl(s):
    s = s.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
    s = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', s)
    s = re.sub(r'\*([^*]+)\*',     r'<i>\1</i>', s)
    return s

def _callout_table(label, text, bg, lcolor):
    """Build a labeled callout box as a ReportLab Table flowable."""
    _reg_fonts()
    txt_color = white if bg == COVER_BG else BODY_TXT
    lb_style = ParagraphStyle('lb', fontName=FBN(bold=True), fontSize=7.5,
                               textColor=lcolor, spaceAfter=2)
    ct_style = ParagraphStyle('ct', fontName=FBN(), fontSize=10,
                               textColor=txt_color, leading=14)
    rows = [[Paragraph(label.upper(), lb_style)],
            [Paragraph(_rl(text), ct_style)]]
    t = Table(rows, colWidths=[CW - 0.8*cm])
    t.setStyle(TableStyle([
        ('BACKGROUND',    (0,0), (-1,-1), bg),
        ('LEFTPADDING',   (0,0), (-1,-1), 10),
        ('RIGHTPADDING',  (0,0), (-1,-1), 10),
        ('TOPPADDING',    (0,0), (0,0),    8),
        ('TOPPADDING',    (0,1), (-1,-1),  2),
        ('BOTTOMPADDING', (0,-1),(-1,-1),  8),
        ('ROUNDEDCORNERS',(0,0), (-1,-1),  3),
    ]))
    return t


def _parse_story(content):
    """Parse markdown-ish content into ReportLab flowables."""
    _reg_fonts()
    B  = ParagraphStyle('b',  fontName=FBN(),       fontSize=11,   textColor=BODY_TXT,
                        alignment=TA_JUSTIFY, firstLineIndent=18, spaceAfter=5, leading=16.5)
    BF = ParagraphStyle('bf', fontName=FBN(),       fontSize=11,   textColor=BODY_TXT,
                        alignment=TA_JUSTIFY, firstLineIndent=0,  spaceAfter=5, leading=16.5,
                        autoLeading='max')  # drop cap 28pt: dòng đầu tự giãn theo glyph, không đè đoạn trên
    SH = ParagraphStyle('sh', fontName=FDN(bold=True), fontSize=13, textColor=NAVY,
                        alignment=TA_LEFT, spaceBefore=14, spaceAfter=5, leading=17)
    SH2= ParagraphStyle('sh2',fontName=FDN(bold=True), fontSize=11, textColor=NAVY,
                        alignment=TA_LEFT, spaceBefore=10, spaceAfter=4, leading=15)
    QT = ParagraphStyle('qt', fontName=FBN(italic=True), fontSize=10, textColor=SUBTEXT,
                        alignment=TA_JUSTIFY, leftIndent=24, rightIndent=24, spaceAfter=6,
                        leading=14)
    CB = ParagraphStyle('cb', fontName=FBN(), fontSize=10, textColor=BODY_TXT,
                        alignment=TA_JUSTIFY, leftIndent=12, rightIndent=12,
                        spaceBefore=8, spaceAfter=8, leading=14,
                        backColor=CALLOUT_BG)
    CB_I=ParagraphStyle('cbi',fontName=FBN(italic=True),fontSize=10,textColor=SUBTEXT,
                        alignment=TA_JUSTIFY, leftIndent=12, rightIndent=12,
                        spaceBefore=2, spaceAfter=8, leading=14,
                        backColor=CALLOUT_BG)

    story = []; first = True
    for raw in [p.strip() for p in content.split('\n\n') if p.strip()]:
        if raw.startswith('## '):
            story += [Spacer(1,4), Paragraph(_xe(raw[3:]).upper(), SH),
                      HRFlowable(width='100%', thickness=0.8, color=GOLD, spaceAfter=5)]
            continue  # drop cap CHỈ cho đoạn đầu chương (đúng SKILL.md) — không reset sau heading
        if raw.startswith('# '):
            story.append(Paragraph(_xe(raw[2:]), SH2)); continue
        if raw.startswith('> '):
            story.append(Paragraph(f'<i>{_xe(raw[2:].strip())}</i>', QT)); continue
        # Labeled callout boxes (!b, !t, !d, !l, !p, !c, !g)
        matched_callout = False
        for pfx, (lbl, bg, lc) in CALLOUT_DEFS.items():
            if raw.startswith(pfx + ' '):
                story.append(Spacer(1, 4))
                story.append(_callout_table(lbl, raw[len(pfx)+1:].strip(), bg, lc))
                story.append(Spacer(1, 4))
                matched_callout = True; break
        if matched_callout: continue
        # Generic callout (no label)
        if raw.startswith('!! '):
            story.append(Paragraph(_rl(raw[3:].strip()), CB)); continue
        if raw.startswith('!i '):
            story.append(Paragraph(f'<i>{_rl(raw[3:].strip())}</i>', CB_I)); continue
        txt = _rl(raw)
        if first:
            # Only do drop cap if text starts with plain char (no HTML tags)
            if txt and txt[0] not in '<&':
                drop = (f'<font name="{FDN(bold=True)}" size="28" color="{_T["accent"]}">'
                        f'{_xe(txt[0])}</font>{txt[1:]}')
                story.append(Paragraph(drop, BF))
            else:
                story.append(Paragraph(txt, BF))
            first = False
        else:
            story.append(Paragraph(txt, B))
    return story

def _p_chapter_content(c, ch, book, start_pg):
    """Flow chapter body text across as many pages as needed."""
    _reg_fonts()
    content = ch.get('content', '')
    if not content: return 0
    story = _parse_story(content)
    if ch.get('section_title'):
        SH = ParagraphStyle('sh0', fontName=FDN(bold=True), fontSize=11, textColor=NAVY,
                            spaceBefore=4, spaceAfter=6)
        story = [Paragraph(ch['section_title'].upper(), SH),
                 HRFlowable(width='100%', thickness=0.8, color=GOLD, spaceAfter=8)] + story

    series = book.get('series_label', book.get('title', ''))
    pg = start_pg; used = 0; first_page = True

    while story:
        _setup_content_page(c, pg, series)
        if first_page:
            _draw_chapter_heading(c, ch)
            Frame(ML, MB, CW, CH - HEADING_H, showBoundary=0).addFromList(story, c)
            first_page = False
        else:
            Frame(ML, MB, CW, CH, showBoundary=0).addFromList(story, c)
        if story: c.showPage(); pg += 1; used += 1

    c.showPage(); used += 1
    return used

# ── PDF builder ────────────────────────────────────────────────────────────────
def build_pdf(book, path):
    _reg_fonts()
    c = pdf_canvas.Canvas(path, pagesize=A4)
    c.setTitle(book.get('title','')); c.setAuthor(book.get('author',''))
    pg = [1]
    def np(): c.showPage(); pg[0] += 1

    _p_cover(c, book); np()

    series = book.get('series_label', book.get('title', ''))

    if book.get('about_author') or book.get('credits'):
        _p_about(c, book, pg[0]); np()

    _p_toc(c, book, pg[0]); np()

    for part in book.get('parts', []):
        _p_part_splash(c, part); np()
        for ch in part.get('chapters', []):
            _p_chapter_opener(c, ch); np()
            used = _p_chapter_content(c, ch, book, pg[0])
            pg[0] += used

    c.save()

# ── DOCX helpers ───────────────────────────────────────────────────────────────
CG_FONT  = 'Cormorant Garamond'
TNR_FONT = 'Times New Roman'

def _dr(p, t, name=TNR_FONT, sz=11, bold=False, italic=False, color=None):
    r = p.add_run(t); r.font.name = name; r.font.size = Pt(sz)
    r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = color
    return r

def _dp(doc, t='', align=WD_ALIGN_PARAGRAPH.LEFT, sz=11, bold=False, italic=False,
        color=None, sb=0, sa=6, fi=0, font=TNR_FONT, ls=None):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    if fi: p.paragraph_format.first_line_indent = DCm(fi)
    if ls:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = ls
    if t: _dr(p, t, name=font, sz=sz, bold=bold, italic=italic, color=color)
    return p

def _set_ls(p, multiple=1.2):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = multiple

def _sp(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

def _drich(p, text, font=TNR_FONT):
    for tok in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text):
        if tok.startswith('**') and tok.endswith('**'):
            _dr(p, tok[2:-2], name=font, bold=True)
        elif tok.startswith('*') and tok.endswith('*'):
            _dr(p, tok[1:-1], name=font, italic=True)
        else:
            _dr(p, tok, name=font)

def _add_docx_page_border(sec, hex_color=None, sz=6, space=20):
    if hex_color is None: hex_color = _T['accent'][1:]
    """Add thin gold border to every page in the DOCX section."""
    pPr = sec._sectPr
    # Remove existing pgBorders if any
    for old in pPr.findall(qn('w:pgBorders')):
        pPr.remove(old)
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(sz))
        border.set(qn('w:space'), str(space))
        border.set(qn('w:color'), hex_color)
        pgBorders.append(border)
    pPr.append(pgBorders)

def _docx_bg_anchor(paragraph, image_path):
    """Insert full-page behind-text background image."""
    W = int(21.0/2.54*914400); H = int(29.7/2.54*914400)
    run = paragraph.add_run()
    run.add_picture(image_path, width=Emu(W), height=Emu(H))
    w_drawing = run._r.find(qn('w:drawing'))
    if w_drawing is None: return
    wp_inline = w_drawing.find(qn('wp:inline'))
    if wp_inline is None: return
    a_blip = wp_inline.find('.//' + qn('a:blip'))
    if a_blip is None: return
    r_embed = a_blip.get(qn('r:embed'))
    if not r_embed: return
    anc = (f'<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
           f' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
           f' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
           f' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
           f' distT="0" distB="0" distL="0" distR="0" simplePos="0"'
           f' relativeHeight="251658240" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="0">'
           f'<wp:simplePos x="0" y="0"/>'
           f'<wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>'
           f'<wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>'
           f'<wp:extent cx="{W}" cy="{H}"/><wp:effectExtent l="0" t="0" r="0" b="0"/>'
           f'<wp:wrapNone/><wp:docPr id="900" name="CoverBg"/><wp:cNvGraphicFramePr/>'
           f'<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
           f'<pic:pic><pic:nvPicPr><pic:cNvPr id="901" name="CoverBg"/>'
           f'<pic:cNvPicPr preferRelativeResize="0"/></pic:nvPicPr>'
           f'<pic:blipFill><a:blip r:embed="{r_embed}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
           f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{W}" cy="{H}"/></a:xfrm>'
           f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
           f'</pic:pic></a:graphicData></a:graphic></wp:anchor>')
    w_drawing.remove(wp_inline)
    w_drawing.append(etree.fromstring(anc))

# ── DOCX builder ───────────────────────────────────────────────────────────────
def build_docx(book, path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width=DCm(21); sec.page_height=DCm(29.7)
    sec.left_margin=DCm(2.5); sec.right_margin=DCm(2.0)
    sec.top_margin=DCm(2.4); sec.bottom_margin=DCm(2.4)

    # Page border (gold, thin)
    _add_docx_page_border(sec, sz=6, space=20)

    cover_img = book.get('cover_image') or str(IMAGES_DIR / 'cover_bg.jpg')

    # ── Cover ──────────────────────────────────────────────────────────────
    if os.path.exists(cover_img):
        p_bg = doc.add_paragraph()
        p_bg.paragraph_format.space_before = Pt(0)
        p_bg.paragraph_format.space_after  = Pt(0)
        _docx_bg_anchor(p_bg, cover_img)
        p_sp = doc.add_paragraph()
        p_sp.paragraph_format.space_before = Pt(0)
        p_sp.paragraph_format.space_after  = Pt(391)
    else:
        _sp(doc, 12)

    # Series label
    if book.get('series_label'):
        p = _dp(doc, book['series_label'].upper(), WD_ALIGN_PARAGRAPH.CENTER,
                sz=8, color=GOLD_RGB, sa=4, font=TNR_FONT)

    # Title
    p_title = doc.add_paragraph(); p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(4)
    r = p_title.add_run(book['title'].upper())
    r.font.name = CG_FONT; r.font.size = Pt(44)
    r.font.bold = True; r.font.color.rgb = WHITE_RGB if os.path.exists(cover_img) else NAVY_RGB

    # Subtitle
    if book.get('subtitle'):
        p_sub = _dp(doc, book['subtitle'], WD_ALIGN_PARAGRAPH.CENTER,
                    sz=11, italic=True, color=CREAM_RGB, sa=12, font=CG_FONT)

    # Tags
    if book.get('tags'):
        tag_str = '  ·  '.join(t.upper() for t in book['tags'])
        _dp(doc, tag_str, WD_ALIGN_PARAGRAPH.CENTER, sz=8, color=LIGHT_RGB, sa=4)

    # Author
    _sp(doc, 2)
    r_auth = doc.add_paragraph(); r_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_auth.paragraph_format.space_before = Pt(0); r_auth.paragraph_format.space_after = Pt(2)
    ra = r_auth.add_run(book.get('author','').upper())
    ra.font.name = CG_FONT; ra.font.size = Pt(14)
    ra.font.bold = True; ra.font.color.rgb = WHITE_RGB if os.path.exists(cover_img) else NAVY_RGB

    if book.get('level_badge'):
        _dp(doc, book['level_badge'].upper(), WD_ALIGN_PARAGRAPH.CENTER,
            sz=8, color=GOLD_RGB, sa=2)
    doc.add_page_break()

    # ── About ──────────────────────────────────────────────────────────────
    if book.get('about_author') or book.get('credits'):
        about_title = book.get('about_title', 'VỀ TÁC GIẢ')
        _dp(doc, about_title, sz=22, bold=True, color=NAVY_RGB, sa=4, font=CG_FONT)
        if book.get('about_subtitle'):
            _dp(doc, book['about_subtitle'], sz=11, italic=True, color=SUBTEXT_RGB,
                sa=8, font=CG_FONT)
        if book.get('author'):
            _dp(doc, book['author'], sz=13, bold=True, color=GOLD_RGB, font=CG_FONT)
        if book.get('author_title'):
            _dp(doc, book['author_title'], sz=9.5, italic=True, color=LIGHT_RGB, sa=2)
        if book.get('author_contact'):
            _dp(doc, book['author_contact'], sz=9, color=LIGHT_RGB, sa=6)
        if book.get('about_author'):
            p = _dp(doc, book['about_author'], sz=11, sa=6, ls=1.2)
            p.paragraph_format.left_indent  = DCm(0.3)
            p.paragraph_format.right_indent = DCm(0.3)
        if book.get('credits'):
            _sp(doc, 1)
            for k, v in book['credits'].items():
                _dp(doc, f"{k.upper()}: {v}", sz=9.5, sa=2)
        doc.add_page_break()

    # ── TOC ────────────────────────────────────────────────────────────────
    _dp(doc, 'MỤC LỤC', sz=22, bold=True, color=NAVY_RGB, sa=6, font=CG_FONT)
    pg_est = 9
    for part in book.get('parts', []):
        _dp(doc, f"PHẦN {part['number']}  ·  {part.get('title','').upper()}",
            sz=9, bold=True, color=GOLD_RGB, sb=12, sa=4, font=CG_FONT)
        for ch in part.get('chapters', []):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent  = DCm(0.4)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            _set_ls(p, 1.15)
            toc_num = (f'{ch["number"]:02d}' if isinstance(ch.get('number'), int)
                       else str(ch.get('number','')))
            _dr(p, f"{toc_num}  ", name=CG_FONT, sz=22, bold=True, color=GOLD_RGB)
            _dr(p, ch['title'], name=CG_FONT, sz=11, bold=True, color=NAVY_RGB)
            _dr(p, f"\t{pg_est}", sz=10, color=LIGHT_RGB)
            if ch.get('subtitle'):
                s = doc.add_paragraph()
                s.paragraph_format.left_indent = DCm(1.8)
                s.paragraph_format.space_before = Pt(0); s.paragraph_format.space_after = Pt(6)
                _set_ls(s, 1.1)
                _dr(s, ch['subtitle'], sz=9, italic=True, color=LIGHT_RGB)
            pg_est += 15
        _sp(doc, 1)
    doc.add_page_break()

    # ── Parts & Chapters ───────────────────────────────────────────────────
    for part in book.get('parts', []):
        _sp(doc, 10)
        _dp(doc, f"PHẦN {part['number']}", WD_ALIGN_PARAGRAPH.CENTER,
            sz=10, bold=True, color=GOLD_RGB, sa=4, font=CG_FONT)
        _dp(doc, part.get('title','').upper(), WD_ALIGN_PARAGRAPH.CENTER,
            sz=28, bold=True, color=NAVY_RGB, sa=6, font=CG_FONT)
        if part.get('subtitle'):
            _dp(doc, part['subtitle'], WD_ALIGN_PARAGRAPH.CENTER,
                sz=12, italic=True, color=SUBTEXT_RGB, font=CG_FONT)
        doc.add_page_break()

        for ch in part.get('chapters', []):
            # Chapter opener (splash style in DOCX: just text on cream bg)
            _sp(doc, 4)
            _dp(doc, 'CHƯƠNG',
                sz=9, italic=True, color=GOLD_RGB, sa=4, font=CG_FONT)
            p_num = doc.add_paragraph()
            p_num.paragraph_format.space_before = Pt(0)
            p_num.paragraph_format.space_after  = Pt(4)
            ch_num_str = (f'{ch["number"]:02d}' if isinstance(ch.get('number'), int)
                          else str(ch.get('number','')))
            r_n = p_num.add_run(ch_num_str)
            r_n.font.name = CG_FONT; r_n.font.size = Pt(80)
            r_n.font.bold = True; r_n.font.color.rgb = GOLD_RGB
            _dp(doc, ch['title'], sz=22, bold=True, color=NAVY_RGB, sa=4, font=CG_FONT)
            if ch.get('subtitle'):
                _dp(doc, ch['subtitle'], sz=11, italic=True, color=SUBTEXT_RGB, sa=6, font=CG_FONT)
            if ch.get('quote'):
                _sp(doc, 1)
                q = doc.add_paragraph()
                q.alignment = WD_ALIGN_PARAGRAPH.CENTER
                q.paragraph_format.left_indent  = DCm(1.5)
                q.paragraph_format.right_indent = DCm(1.5)
                q.paragraph_format.space_after  = Pt(2)
                _dr(q, f'“{ch["quote"]}”', sz=10, italic=True, color=SUBTEXT_RGB)
                if ch.get('quote_attribution'):
                    qa = doc.add_paragraph()
                    qa.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    qa.paragraph_format.space_after = Pt(4)
                    _dr(qa, f'— {ch["quote_attribution"]}',
                        name=CG_FONT, sz=8.5, bold=True, color=GOLD_RGB)
            doc.add_page_break()

            # Chapter content
            if ch.get('section_title'):
                _dp(doc, ch['section_title'].upper(),
                    sz=12, bold=True, color=NAVY_RGB, sb=4, sa=8, font=CG_FONT)

            first = True
            for raw in [p.strip() for p in ch.get('content','').split('\n\n') if p.strip()]:
                if raw.startswith('## '):
                    _dp(doc, raw[3:].upper(), sz=12, bold=True, color=NAVY_RGB,
                        sb=12, sa=6, font=CG_FONT)
                    first = True; continue
                if raw.startswith('# '):
                    _dp(doc, raw[2:], sz=11, bold=True, color=NAVY_RGB,
                        sb=8, sa=4, font=CG_FONT)
                    first = True; continue
                if raw.startswith('> '):
                    q = doc.add_paragraph()
                    q.paragraph_format.left_indent  = DCm(1.5)
                    q.paragraph_format.right_indent = DCm(1.5)
                    q.paragraph_format.space_after  = Pt(6)
                    _set_ls(q, 1.15)
                    _dr(q, raw[2:].strip(), sz=10, italic=True, color=SUBTEXT_RGB)
                    continue
                # Labeled callout boxes for DOCX
                matched_lbl = False
                for pfx in CALLOUT_DEFS:
                    if raw.startswith(pfx + ' '):
                        lbl_text, _, _ = CALLOUT_DEFS[pfx]
                        fill_hex = CALLOUT_DOCX_FILLS[pfx]
                        lbl_hex  = CALLOUT_DOCX_LABEL_COLORS[pfx]
                        # Label paragraph
                        lp = doc.add_paragraph()
                        lp.paragraph_format.left_indent  = DCm(0.5)
                        lp.paragraph_format.right_indent = DCm(0.5)
                        lp.paragraph_format.space_before = Pt(8)
                        lp.paragraph_format.space_after  = Pt(0)
                        lp_r = lp.add_run(lbl_text.upper())
                        lp_r.font.name = TNR_FONT; lp_r.font.size = Pt(7.5)
                        lp_r.font.bold = True
                        lp_r.font.color.rgb = RGBColor(
                            int(lbl_hex[0:2],16), int(lbl_hex[2:4],16), int(lbl_hex[4:6],16))
                        for pp in [lp]:
                            pPr = pp._p.get_or_add_pPr()
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
                            shd.set(qn('w:fill'), fill_hex); pPr.append(shd)
                        # Body paragraph
                        q = doc.add_paragraph()
                        q.paragraph_format.left_indent  = DCm(0.5)
                        q.paragraph_format.right_indent = DCm(0.5)
                        q.paragraph_format.space_before = Pt(2)
                        q.paragraph_format.space_after  = Pt(8)
                        _set_ls(q, 1.15)
                        _drich(q, raw[len(pfx)+1:].strip())
                        pPr2 = q._p.get_or_add_pPr()
                        shd2 = OxmlElement('w:shd')
                        shd2.set(qn('w:val'), 'clear'); shd2.set(qn('w:color'), 'auto')
                        shd2.set(qn('w:fill'), fill_hex); pPr2.append(shd2)
                        matched_lbl = True; break
                if matched_lbl: continue
                if raw.startswith('!! ') or raw.startswith('!i '):
                    q = doc.add_paragraph()
                    q.paragraph_format.left_indent  = DCm(0.5)
                    q.paragraph_format.right_indent = DCm(0.5)
                    q.paragraph_format.space_before = Pt(6)
                    q.paragraph_format.space_after  = Pt(6)
                    _set_ls(q, 1.15)
                    _drich(q, raw[3:].strip())
                    # Light shading
                    pPr = q._p.get_or_add_pPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'EEECE4')
                    pPr.append(shd)
                    continue
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(6)
                _set_ls(p, 1.2)
                if first:
                    p.paragraph_format.first_line_indent = DCm(0)
                    first = False
                else:
                    p.paragraph_format.first_line_indent = DCm(0.7)
                _drich(p, raw)

    doc.save(path)

# ── Entry point ────────────────────────────────────────────────────────────────
def main():
    args = sys.argv[1:]
    docx_only = '--docx-only' in args
    pdf_only  = '--pdf-only'  in args
    args = [a for a in args if not a.startswith('--')]

    if not args:
        print('Usage: python generate_book.py <book_data.json> [output_dir] [--docx-only] [--pdf-only]')
        print(f'Script: {SCRIPT_DIR}')
        print(f'Assets: {ASSETS_DIR}')
        sys.exit(1)

    _reg_fonts()
    json_path  = args[0]
    output_dir = args[1] if len(args) > 1 else str(Path(json_path).parent)
    os.makedirs(output_dir, exist_ok=True)

    book = json.load(open(json_path, encoding='utf-8'))
    ts   = datetime.now().strftime('%Y%m%d_%H%M%S')
    base = slug(book.get('title', 'book'))

    if not pdf_only:
        dp = os.path.join(output_dir, f'{base}_{ts}.docx')
        print('Generating DOCX...')
        build_docx(book, dp)
        print(f'  DOCX: {dp}')

    if not docx_only:
        pp = os.path.join(output_dir, f'{base}_{ts}.pdf')
        print('Generating PDF...')
        build_pdf(book, pp)
        print(f'  PDF : {pp}')

    print('Done.')

if __name__ == '__main__':
    main()
