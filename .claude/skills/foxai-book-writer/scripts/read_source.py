#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
read_source.py — Extract structured text from source documents for translation
Supports: .docx, .txt, .md, .pdf

Usage:
    python read_source.py <file_path>
    python read_source.py <file_path> --preview   # chỉ hiện 10 block đầu
    python read_source.py <file_path> --plain     # plain text (không JSON)

Output (JSON):
    {
      "source_file": "...",
      "detected_format": "docx|txt|pdf",
      "total_blocks": 123,
      "blocks": [
        {"type": "h1|h2|h3|p|quote|list", "text": "..."},
        ...
      ]
    }
"""

import sys, json, os, re
from pathlib import Path


def _clean(s):
    return re.sub(r'\s+', ' ', s).strip()


def read_docx(path):
    """Extract structured blocks from DOCX preserving heading hierarchy."""
    from docx import Document
    doc = Document(path)
    blocks = []
    for para in doc.paragraphs:
        t = _clean(para.text)
        if not t:
            continue
        sn = para.style.name.lower() if para.style else ''
        if 'heading 1' in sn or sn.startswith('heading1'):
            blocks.append({'type': 'h1', 'text': t})
        elif 'heading 2' in sn or sn.startswith('heading2'):
            blocks.append({'type': 'h2', 'text': t})
        elif 'heading 3' in sn or sn.startswith('heading3'):
            blocks.append({'type': 'h3', 'text': t})
        elif 'quote' in sn or 'block' in sn:
            blocks.append({'type': 'quote', 'text': t})
        elif para.paragraph_format.left_indent and para.paragraph_format.left_indent.pt > 0:
            blocks.append({'type': 'list', 'text': t})
        else:
            blocks.append({'type': 'p', 'text': t})
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(_clean(c.text) for c in row.cells if c.text.strip())
            if row_text:
                blocks.append({'type': 'table_row', 'text': row_text})
    return blocks


def read_txt(path):
    """Extract blocks from plain text / markdown."""
    with open(path, encoding='utf-8', errors='replace') as f:
        content = f.read()
    blocks = []
    for line in content.split('\n'):
        t = line.rstrip()
        if not t.strip():
            continue
        if t.startswith('### '):
            blocks.append({'type': 'h3', 'text': _clean(t[4:])})
        elif t.startswith('## '):
            blocks.append({'type': 'h2', 'text': _clean(t[3:])})
        elif t.startswith('# '):
            blocks.append({'type': 'h1', 'text': _clean(t[2:])})
        elif t.startswith('> '):
            blocks.append({'type': 'quote', 'text': _clean(t[2:])})
        elif re.match(r'^[-*+]\s', t):
            blocks.append({'type': 'list', 'text': _clean(t[2:])})
        else:
            blocks.append({'type': 'p', 'text': _clean(t)})
    return blocks


def read_pdf(path):
    """Extract text from PDF — tries pdftotext first, then PyMuPDF."""
    import subprocess

    # Method 1: pdftotext
    try:
        result = subprocess.run(
            ['pdftotext', '-layout', '-enc', 'UTF-8', path, '-'],
            capture_output=True, timeout=60
        )
        if result.returncode == 0:
            text = result.stdout.decode('utf-8', errors='replace')
            blocks = []
            cur_para = []
            for line in text.split('\n'):
                stripped = line.strip()
                if not stripped:
                    if cur_para:
                        blocks.append({'type': 'p', 'text': ' '.join(cur_para)})
                        cur_para = []
                else:
                    # Simple heuristic: ALL CAPS short line → heading
                    if len(stripped) < 80 and stripped.isupper() and len(stripped.split()) <= 8:
                        if cur_para:
                            blocks.append({'type': 'p', 'text': ' '.join(cur_para)})
                            cur_para = []
                        blocks.append({'type': 'h2', 'text': stripped})
                    else:
                        cur_para.append(stripped)
            if cur_para:
                blocks.append({'type': 'p', 'text': ' '.join(cur_para)})
            return blocks
    except FileNotFoundError:
        pass  # pdftotext not installed
    except Exception as e:
        print(f'[pdftotext error: {e}]', file=sys.stderr)

    # Method 2: PyMuPDF (fitz)
    try:
        import fitz
        doc = fitz.open(path)
        blocks = []
        for pg in doc:
            page_dict = pg.get_text('dict')
            for block in page_dict.get('blocks', []):
                if block.get('type') == 0:  # text block
                    lines = []
                    for line in block.get('lines', []):
                        span_text = ' '.join(s['text'] for s in line.get('spans', []))
                        span_text = _clean(span_text)
                        if span_text:
                            # Detect heading by font size
                            sizes = [s.get('size', 11) for s in line.get('spans', []) if s.get('text')]
                            avg_size = sum(sizes) / len(sizes) if sizes else 11
                            if avg_size > 15:
                                if lines:
                                    blocks.append({'type': 'p', 'text': ' '.join(lines)})
                                    lines = []
                                blocks.append({'type': 'h2' if avg_size > 18 else 'h3',
                                               'text': span_text})
                            else:
                                lines.append(span_text)
                    if lines:
                        blocks.append({'type': 'p', 'text': ' '.join(lines)})
        return blocks
    except ImportError:
        pass
    except Exception as e:
        print(f'[PyMuPDF error: {e}]', file=sys.stderr)

    return [{'type': 'p', 'text':
             'LỖI: Không đọc được PDF. Cần cài pdftotext (poppler-utils) hoặc PyMuPDF (pip install pymupdf).'}]


def estimate_language(blocks):
    """Rough language detection from first ~2000 chars of text."""
    sample = ' '.join(b['text'] for b in blocks[:30])[:2000]
    # Vietnamese markers
    vi_chars = len(re.findall(r'[àáạảãăắặẳẵâấậẩẫèéẹẻẽêếệểễìíịỉĩòóọỏõôốộổỗơớợởỡùúụủũưứựửữỳýỵỷỹđ]', sample, re.I))
    # Chinese/Japanese markers
    cjk = len(re.findall(r'[一-鿿぀-ゟ゠-ヿ]', sample))
    # French markers
    fr_words = len(re.findall(r'\b(le|la|les|des|du|de|et|est|dans|pour|avec|sur|par|ce|qui|que)\b', sample, re.I))
    # English markers
    en_words = len(re.findall(r'\b(the|and|is|are|was|were|this|that|with|for|from|have|has|been)\b', sample, re.I))

    if vi_chars > 10: return 'vi'
    if cjk > 30:
        # distinguish Chinese vs Japanese
        kana = len(re.findall(r'[぀-ゟ゠-ヿ]', sample))
        return 'ja' if kana > 5 else 'zh'
    if fr_words > en_words * 0.6 and fr_words > 5: return 'fr'
    if en_words > 5: return 'en'
    return 'unknown'


def main():
    args = sys.argv[1:]
    if not args:
        print('Usage: python read_source.py <file> [--preview] [--plain]')
        sys.exit(1)

    path = args[0]
    preview_only = '--preview' in args
    plain_mode   = '--plain'   in args

    if not os.path.exists(path):
        print(f'File not found: {path}', file=sys.stderr); sys.exit(1)

    ext = Path(path).suffix.lower()
    fmt = 'unknown'

    if ext == '.docx':
        blocks = read_docx(path); fmt = 'docx'
    elif ext in ('.txt', '.md'):
        blocks = read_txt(path);  fmt = 'txt'
    elif ext == '.pdf':
        blocks = read_pdf(path);  fmt = 'pdf'
    else:
        print(f'Unsupported format: {ext}. Supported: .docx .txt .md .pdf', file=sys.stderr)
        sys.exit(1)

    lang = estimate_language(blocks)

    if plain_mode:
        for b in (blocks[:10] if preview_only else blocks):
            prefix = {'h1':'# ','h2':'## ','h3':'### ','quote':'> ','list':'• '}.get(b['type'],'')
            print(prefix + b['text'])
        return

    out = {
        'source_file':     path,
        'detected_format': fmt,
        'detected_language': lang,
        'total_blocks':    len(blocks),
    }
    if preview_only:
        out['blocks'] = blocks[:10]
        out['note'] = f'Showing 10/{len(blocks)} blocks. Run without --preview for full output.'
    else:
        out['blocks'] = blocks

    # Ensure stdout is UTF-8
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    print(json.dumps(out, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()
